"""
PDF / DOCX fayllaridan sof matn (string) ajratish.
Gemini'ga butun fayl emas, shu matn yuboriladi.
"""
import os


def extract_text(file_field) -> str:
    """Django FileField (PDF yoki DOCX) → sof matn."""
    if not file_field:
        return ''
    name = file_field.name.lower()
    try:
        if name.endswith('.pdf'):
            return _from_pdf(file_field)
        if name.endswith('.docx'):
            return _from_docx(file_field)
    except Exception:
        return ''
    return ''


def _from_pdf(file_field) -> str:
    import fitz  # PyMuPDF
    file_field.open('rb')
    raw = file_field.read()
    file_field.close()
    doc = fitz.open(stream=raw, filetype='pdf')
    parts = []
    for page in doc:
        parts.append(page.get_text('text'))
    doc.close()
    return '\n'.join(parts).strip()


def _from_docx(file_field) -> str:
    import docx  # python-docx
    file_field.open('rb')
    document = docx.Document(file_field)
    file_field.close()
    lines = [p.text for p in document.paragraphs if p.text.strip()]
    # Jadvallardagi matn ham
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return '\n'.join(lines).strip()
